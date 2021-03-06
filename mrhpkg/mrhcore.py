#!/usr/bin/env python
# -*- coding: utf-8 -*-
# Filename: mrhcore.py
# Usage: Medical Review Helper Core Library

import configparser
import csv
import datetime
import os
import pickle
import threading
import time

from PyQt5.QtCore import QThread, pyqtSignal
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

import mrhpkg.mrhimp as mrhimp
import mrhpkg.mrhini as mrhini
import mrhpkg.spider.cnki as cnkispider
import mrhpkg.spider.pubmed as pmdspider
import mrhpkg.spider.wanfang as wfspider


class MrhConfig:
    """Parse mrhelper.ini file.
    Public:
        save: save mrhelper.ini file

    Returns:
        ini: all .ini field
        sci: sci.csv reader
        hexin: hexin.csv reader
        info: Text Template for Reference Info
    """

    def __init__(self):
        self.inifile = 'mrhelper.ini'
        self.ini = self._get_ini()
        self.sci = self._get_sci()
        self.hexin = self._get_hexin()
        self.info = self._get_info()

    def _get_ini(self):
        config = configparser.ConfigParser()
        if not os.path.exists(self.inifile):
            with open(self.inifile, 'w', encoding='utf-8') as inifile:
                inifile.write(mrhini.INI)
        try:
            config.read(self.inifile, encoding='utf-8')
        except configparser.MissingSectionHeaderError:
            config.read(self.inifile, encoding='utf-8-sig')
        return config

    def _get_sci(self):
        sci = {}
        with open(self.ini['Resource']['sci'], newline='', encoding='utf-8') as scifile:
            csvreader = csv.DictReader(scifile)
            for row in csvreader:
                sci[row['title'].upper()] = row['quarter']
        return sci

    def _get_hexin(self):
        hexin = set()
        with open(self.ini['Resource']['hexin'], newline='', encoding='utf-8') as hexinfile:
            csvreader = csv.DictReader(hexinfile)
            for row in csvreader:
                hexin.add(row['journal'])
        return hexin

    def _get_info(self):
        with open(self.ini['Resource']['info'], 'r', encoding='utf-8') as infofile:
            info = infofile.read()
        return info

    @staticmethod
    def save(config):
        """save mrhelper.ini file."""
        with open('mrhelper.ini', 'w', encoding='utf-8') as inifile:
            config.ini.write(inifile)


class MrhIo(QThread):
    """Manage Data Processing."""

    sigmsg = pyqtSignal(str)
    sigmrh = pyqtSignal(object)
    sigover = pyqtSignal()

    def __init__(self, mrhproject, filepaths, mode):
        super().__init__()
        self.mrhproject = mrhproject
        self.filepaths = filepaths
        self.mode = mode
        self.iothread = ''
        self.checkiothread = ''

    def run(self):
        """Run Function For QThread."""

        if self.mode == 'save':
            self.iothread = threading.Thread(target=self.mrhsave)
        elif self.mode == 'open':
            self.iothread = threading.Thread(target=self.mrhopen)
        elif self.mode == 'add':
            self.iothread = threading.Thread(target=self.mrhadd)
        else:
            self.sigmsg.emit('MrhIo Mode Needed...')

        if self.iothread:
            self.iothread.start()
            self.checkiothread = threading.Thread(target=self._check_iothread)
            self.checkiothread.start()

    def _check_iothread(self):
        start_time = time.time()
        while self.iothread.is_alive():
            self.msleep(100)
            elapse_time = str(round(time.time() - start_time, 1))
            self.sigmsg.emit(
                f'{self.mode.upper()} Time Elapsed: {elapse_time} seconds')
        if self.mode == 'open' or self.mode == 'add':
            self.sigmrh.emit(self.mrhproject)
        self.sigmsg.emit(f'{self.mode.upper()} is DONE!')

    def mrhsave(self):
        """Save Mrhproject."""
        with open(self.filepaths, 'wb') as savefile:
            pickle.dump(self.mrhproject, savefile)

    def mrhopen(self):
        """Open Mrhproject."""
        with open(self.filepaths, 'rb') as openfile:
            self.mrhproject = pickle.load(openfile)

    def mrhadd(self):
        """Add exported files from supported databases."""
        for filepath in self.filepaths:
            data = mrhimp.getdata(filepath)
            baserid = len(self.mrhproject.mrhdata)
            if data['mrhdata'] != -1:
                for mrhitem in data['mrhdata']:
                    mrhitem.rid += baserid
                self.mrhproject.mrhdata += data['mrhdata']
                self.mrhproject.rawdata += data['rawdata']


class MrhTable:
    """Manage Tab_READ datatable."""

    @staticmethod
    def markitem(config, mrhitem):
        itemcolor = {}
        itemcolor['year'] = MrhTable._mark_year(config, mrhitem)
        itemcolor['type'] = MrhTable._mark_type(config, mrhitem)
        itemcolor['database'] = MrhTable._mark_database(config, mrhitem)
        itemcolor['journal'] = MrhTable._mark_journal(config, mrhitem)
        return itemcolor

    @staticmethod
    def _mark_year(config, mrhitem):
        now_year = datetime.datetime.now().year
        return 'lightgreen' if mrhitem.year and now_year - int(mrhitem.year) < 6 else ''

    @staticmethod
    def _mark_type(config, mrhitem):
        if mrhitem.type == 'Review':
            return 'lightgreen'
        else:
            if mrhitem.type == 'Article' or mrhitem.type == 'Journal Article' or 'Journal Article' in mrhitem.type[0]:
                return 'gold'

    @staticmethod
    def _mark_database(config, mrhitem):
        return 'lightgreen' if mrhitem.pmcid else ''

    @staticmethod
    def _mark_journal(config, mrhitem):
        if mrhitem.journal:
            if mrhitem.database == 'WOS' or mrhitem.database == 'PUBMED':
                if mrhitem.journal in config.sci:
                    quarter = config.sci[mrhitem.journal]
                    if quarter == 'Q1':
                        return 'lightgreen'
                    elif quarter == 'Q2':
                        return 'gold'
                    elif quarter == 'Q3':
                        return 'yellow'
                    elif quarter == 'Q4':
                        return 'lightyellow'
                    else:                        
                        return 'blue'
            else:
                if mrhitem.journal in config.hexin:
                    return 'lightgreen'

    @staticmethod
    def check_viewoptions(viewoptions, currentrid, currentitem, mrhitem):
        """Check Mrhitem Displayed."""
        relate = MrhTable._check_relate(
            viewoptions, currentrid, currentitem, mrhitem)
        abstract = MrhTable._check_abstract(viewoptions, mrhitem)
        fiveyears = MrhTable._check_fiveyears(viewoptions, mrhitem)
        reftype = MrhTable._check_type(viewoptions, mrhitem)
        use = MrhTable._check_use(viewoptions, mrhitem)
        search = MrhTable._check_search(viewoptions, mrhitem)
        result = abstract and fiveyears and reftype and use and search and relate
        return result

    @staticmethod
    def _check_relate(viewoptions, currentrid, currentitem, mrhitem):
        if viewoptions['relate'] == 2:
            if currentrid == -1:
                return True
            else:
                if currentrid == mrhitem.rid:
                    return True
                else:
                    if currentitem.lcs:
                        if mrhitem.rid in currentitem.lcs:
                            return True
                    if currentitem.lcr:
                        if mrhitem.rid in currentitem.lcr:
                            return True
                    if mrhitem.doi:
                        if mrhitem.doi == currentitem.doi:
                            return True
                    if mrhitem.title:
                        if mrhitem.title == currentitem.title:
                            return True
                    return False
        else:
            return True

    @staticmethod
    def _check_abstract(viewoptions, mrhitem):
        option = viewoptions['abstract']
        if option == 2 and mrhitem.abstract:
            return True
        elif option == 0 and not mrhitem.abstract:
            return True
        elif option == 1:
            return True
        else:
            return False

    @staticmethod
    def _check_fiveyears(viewoptions, mrhitem):
        option = viewoptions['fiveyears']
        now_year = datetime.datetime.now().year
        if mrhitem.year:
            if option == 2 and now_year - int(mrhitem.year) < 6:
                return True
            elif option == 0 and now_year - int(mrhitem.year) >= 6:
                return True
            elif option == 1:
                return True
            else:
                return False
        elif option == 1:
            return True
        else:
            return False

    @staticmethod
    def _check_type(viewoptions, mrhitem):
        option = viewoptions['type']
        if option == 2 and (mrhitem.type == 'Review' or mrhitem.type == 'Article' or 'Journal Article' in mrhitem.type):
            return True
        elif option == 0 and mrhitem.type != 'Review' and mrhitem.type != 'Article' \
                and 'Journal Article' not in mrhitem.type:
            return True
        elif option == 1:
            return True
        else:
            return False

    @staticmethod
    def _check_use(viewoptions, mrhitem):
        option = viewoptions['use']
        if option == 2 and mrhitem.use == 2:
            return True
        elif option == 0 and mrhitem.use == 0:
            return True
        elif option == 1:
            return True
        else:
            return False

    @staticmethod
    def _check_search(viewoptions, mrhitem):
        option = viewoptions['search']
        if option:
            result = (option in mrhitem.title) or (option in mrhitem.abstract) or \
                MrhTable._check_search_author(option, mrhitem) or (
                option in mrhitem.journal)
            return result
        else:
            return True

    @staticmethod
    def _check_search_author(option, mrhitem):
        for author in mrhitem.author:
            if option.lower() in author.lower():
                return True
        return False


class MrhWeb(QThread):
    sigmsg = pyqtSignal(str)
    sigmrh = pyqtSignal(object)

    def __init__(self, mrhdata, mrhproject, config):
        super().__init__()
        self.data = mrhdata
        self.mrhproject = mrhproject
        self.config = config
        self.status = {'PUBMED': [0, 0], 'CNKI': [0, 0], 'WANFANG': [0, 0]}
        self.checkwebthread = ''

    def run(self):
        databases = ['PUBMED', 'CNKI', 'WANFANG']
        datadict = {}
        threads = []
        for item in self.data:
            if item.database in databases:
                datadict.setdefault(item.database, []).append(item)
        for database in datadict:
            self.status[database][0] = len(datadict[database])
        for database in databases:
            if datadict.setdefault(database):
                database_thread = threading.Thread(
                    target=self._get_info, args=(datadict[database], database))
                threads.append(database_thread)
                database_thread.start()
        for thread in threads:
            thread.join()
        self.sigmrh.emit(self.mrhproject)

    def _show_progress(self):
        message = ['Retrieve: ']
        for key, value in self.status.items():
            message.append(f"{key.capitalize()} {value[1]}/{value[0]}")
        message = ' '.join(message)
        self.sigmsg.emit(message)

    def _get_info(self, dbdata, database):
        threads = []
        for index, item in enumerate(dbdata):
            if item.database == 'WANFANG' and item.link:
                thread = threading.Thread(target=MrhSpider(
                    item, self.data, self.config).wanfang)
            elif item.database == 'CNKI' and item.link:
                thread = threading.Thread(target=MrhSpider(
                    item, self.data, self.config).cnki)
            elif item.database == 'PUBMED':
                thread = threading.Thread(target=MrhSpider(
                    item, self.data, self.config).pubmed)
            else:
                continue
            thread.name = item.database + '_' + str(index)
            threads.append(thread)
            thread.start()
            count = self._count_thread()
            while count:
                time.sleep(1)
                count = self._count_thread()
            self.status[database][1] = index + 1
            self._show_progress()
        for thread in threads:
            thread.join()

    def _count_thread(self):
        max_pubmed = int(self.config.ini['Pubmed']['threads'])
        max_cnki = int(self.config.ini['CNKI']['threads'])
        max_wanfang = int(self.config.ini['Wanfang']['threads'])
        pubmed = 0
        cnki = 0
        wanfang = 0
        for thread in threading.enumerate():
            name = thread.getName().split('_')[0]
            if name == 'PUBMED':
                pubmed += 1
            elif name == 'CNKI':
                cnki += 1
            elif name == 'WANFANG':
                wanfang += 1
            else:
                pass
        if pubmed > max_pubmed or cnki > max_cnki or wanfang > max_wanfang:
            return True
        return False


class MrhSpider:
    def __init__(self, mrhitem, mrhdata, config):
        self.item = mrhitem
        self.data = mrhdata
        self.config = config

    def wanfang(self):
        wfid = self.item.link.split('id=')[1]
        time_out = int(self.config.ini['Wanfang']['time_out'])
        retries = int(self.config.ini['Wanfang']['retries'])
        proxies = {
            'http': self.config.ini['Proxy']['http'],
            'https': self.config.ini['Proxy']['https']
        }
        reflist = wfspider.getdata(
            wfid, time_out=time_out, retries=retries, proxies=proxies)
        crlist = reflist['cr'][0] if reflist['cr'] != -1 else -1
        cslist = reflist['cs'][0] if reflist['cs'] != -1 else -1
        self.item.cr = reflist['cr'][1] if reflist['cr'] != -1 else -1
        self.item.cs = reflist['cs'][1] if reflist['cs'] != -1 else -1
        crdoi = set()
        crtitle = set()
        if crlist != -1 and crlist:
            for item in crlist:
                if item['DOI']:
                    crdoi.add(item['DOI'])
                if item['Title']:
                    crtitle.add(item['Title'][0])
        csdoi = set()
        cstitle = set()
        if cslist != -1 and cslist:
            for item in cslist:
                if item['DOI']:
                    csdoi.add(item['DOI'])
                if item['Title']:
                    cstitle.add(item['Title'][0])
        self.item.lcr = []
        self.item.lcs = []
        for item in self.data:
            if item.doi in crdoi or item.title in crtitle:
                self.item.lcr.append(item.rid)
            if item.doi in csdoi or item.title in cstitle:
                self.item.lcs.append(item.rid)

    def cnki(self):
        link = self.item.link
        time_out = int(self.config.ini['CNKI']['time_out'])
        retries = int(self.config.ini['CNKI']['retries'])
        proxies = {
            'http': self.config.ini['Proxy']['http'],
            'https': self.config.ini['Proxy']['https']
        }
        count = cnkispider.getcount(
            link, time_out=time_out, retries=retries, proxies=proxies)
        reflist = cnkispider.getdata(
            link, time_out=time_out, retries=retries, proxies=proxies)
        if count:
            self.item.cr = count['REFERENCE']
            self.item.cs = count['CITING']
        if reflist['cr'] == -1 or reflist['cs'] == -1:
            return
        cstitleset = set()
        for item in reflist['cs']:
            soup = BeautifulSoup(item, 'lxml')
            for link in soup.find_all('a'):
                if link.get('target') == 'kcmstarget':
                    citetitle = link.get_text()
                    cstitleset.add(citetitle)
                if link.get('onclick'):
                    if 'OpenCRLDENG' in link.get('onclick'):
                        reftitle = link.get_text()
                        cstitleset.add(reftitle)
        crtitleset = set()
        for item in reflist['cr']:
            soup = BeautifulSoup(item, 'lxml')
            for link in soup.find_all('a'):
                if link.get('target') == 'kcmstarget':
                    reftitle = link.get_text()
                    crtitleset.add(reftitle)
                if link.get('onclick'):
                    if 'OpenCRLDENG' in link.get('onclick'):
                        reftitle = link.get_text()
                        crtitleset.add(reftitle)
        self.item.lcs = []
        self.item.lcr = []
        for item in self.data:
            if item.title in cstitleset:
                self.item.lcs.append(item.rid)
            if item.title in crtitleset:
                self.item.lcr.append(item.rid)

    def pubmed(self):
        api_key = self.config.ini['Pubmed']['api_key']
        if not api_key:
            api_key = None
        time_out = int(self.config.ini['Pubmed']['time_out'])
        retries = int(self.config.ini['Pubmed']['retries'])
        proxies = {
            'http': self.config.ini['Proxy']['http'],
            'https': self.config.ini['Proxy']['https']
        }
        reflist = pmdspider.getdata(self.item.pmid, api_key=api_key, time_out=time_out, retries=retries,
                                    proxies=proxies)
        self.item.cr = len(reflist['cr']) if reflist['cr'] != -1 else -1
        self.item.cs = len(reflist['cs']) if reflist['cs'] != -1 else -1
        self.item.lcr = []
        self.item.lcs = []
        for item in self.data:
            if reflist['cr'] != -1 and reflist['cs'] != -1:
                if item.pmid in reflist['cr']:
                    self.item.lcr.append(item.rid)
                elif item.pmid in reflist['cs']:
                    self.item.lcs.append(item.rid)


class MrhExport:
    def __init__(self, mrhproject, savepath, reftree, config):
        self.mrhproject = mrhproject
        self.savepath = savepath
        self.reftree = reftree
        self.document = Document(config['Resource']['docx'])
        self.endnote()
        self.docx()

    def endnote(self):
        filepath = os.path.join(self.savepath, 'endnote.txt')
        with open(filepath, 'w', encoding='utf-8') as datafile:
            for item in self.mrhproject.mrhdata:
                # Type
                datafile.write('%0 Journal Article\n')
                # author
                for author in item.author:
                    datafile.write('%A ')
                    datafile.write(author)
                    datafile.write('\n')
                # title
                datafile.write('%T ')
                datafile.write(item.title)
                datafile.write('\n')
                # Journal
                datafile.write('%B ')
                datafile.write(item.journal)
                datafile.write('\n')
                # year
                datafile.write('%D ')
                datafile.write(item.year)
                datafile.write('\n')
                # Volumn
                datafile.write('%V ')
                datafile.write(item.volumn)
                datafile.write('\n')
                # Issue
                datafile.write('%N ')
                datafile.write(item.issue)
                datafile.write('\n')
                # Page
                datafile.write('%P ')
                datafile.write(item.page)
                datafile.write('\n')
                # pmid
                datafile.write('%M ')
                datafile.write(item.pmid)
                datafile.write('\n')
                # pmcid
                datafile.write('%2 ')
                datafile.write(item.pmcid)
                datafile.write('\n')
                # DOI
                datafile.write('%R ')
                datafile.write(item.doi)
                datafile.write('\n')
                # Abstract
                datafile.write('%X ')
                datafile.write(item.abstract)
                datafile.write('\n')
                # Label
                datafile.write('%F ')
                datafile.write('mr.'+str(item.rid + 1))
                datafile.write('\n')
                # Alternate Journal
                datafile.write('%O ')
                datafile.write(item.journal_en)
                datafile.write('\n')
                # Translated Author
                for author_en in item.author_en:
                    datafile.write('%H ')
                    datafile.write(author_en)
                    datafile.write('\n')
                # Translated Title
                datafile.write('%Q ')
                datafile.write(item.title_en)
                datafile.write('\n')
                # End of Record
                datafile.write('\n')

    def docx(self):

        #####################################
        # Get Styles
        #####################################
        # styles=document.styles
        # paragraph_styles = [s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH]
        # for style in paragraph_styles:
        #     print(style.name)
        #####################################
        # Add Title
        self.document.add_heading('Review Title', 1)

        # Add Author
        paragraph = self.document.add_paragraph('')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        author = paragraph.add_run('Author List')
        author.bold = True
        author.italic = True
        paragraph.add_run(' ')

        # Add Abstract
        self.document.add_heading('Abstract', 4)
        self.document.add_paragraph('Abstract Content')

        # Add Keywords
        paragraph = self.document.add_paragraph('Keywords: ', 'Heading 4')
        paragraph.add_run('Keywords List')

        # Add Heading&References
        for node in self.reftree:
            if node[2] == -1:
                if node[0] != '- -':
                    heading = self.document.add_paragraph(node[0])
                    if node[1] == 1:
                        heading.style = self.document.styles['Heading 2']
                    if node[1] == 2:
                        heading.style = self.document.styles['Heading 3']
            else:
                item = self.mrhproject.mrhdata[int(node[2])]
                # Add References
                self._docx_addref(item)
        # Add Reference Text
        self.document.add_heading('Reference', 2)
        # Save
        docxfile = os.path.join(self.savepath, 'mrhelper.docx')
        self.document.save(docxfile)

    def _docx_addref(self, item):
        if item.author:
            if isinstance(item.author, list):
                authorname = item.author[0].split(',')[0]
                authorname = authorname.split(' ')[0]
                authorname = authorname.split(';')[0]
            else:
                authorname = item.author
        else:
            authorname = ''
        if item.year:
            year = item.year
        else:
            year = None
        if item.title:
            title = item.title
        else:
            title = ''
        seq = item.rid + 1
        citation = "{%s, %s, %s%d}" % (authorname, year, 'mr.', seq)
        iv = item.iv if item.iv else 'iv'
        dv = item.dv if item.dv else 'dv'
        rel = '+' if item.relation == 2 else '±'
        rel = '-' if item.relation == 0 else rel
        description = item.reftext
        text = "%s%s %s%s%s %s %s" % (
            authorname, citation, iv, rel, dv, title, description)
        self.document.add_paragraph(text)
